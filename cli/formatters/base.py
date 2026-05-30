import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


def add_run_with_font(paragraph, text, font_name='宋体', font_size=None,
                       bold=False, color=None, italic=False, underline=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if underline is not None:
        run.font.underline = underline
    if color:
        run.font.color.rgb = color
    return run


def add_shading(paragraph, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    paragraph._element.get_or_add_pPr().append(shading)


FONT_SANHAO = Pt(16)
FONT_SIHAO = Pt(14)
FONT_XIAOSI = Pt(12)
FONT_WUHAO = Pt(10.5)
RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

COLOR_REF = {"黑色": BLACK, "红色": RED, "蓝色": BLUE, "绿色": RGBColor(0x00, 0x80, 0x00)}
SIZE_REF = {"小三": Pt(15), "四号": Pt(14), "小四": Pt(12), "五号": Pt(10.5), "小五": Pt(9)}


EXAM_ORDER_KEYWORDS = [
    ("第一次月考", 0),
    ("第二次月考", 2),
    ("月考", 0),
    ("期中", 1),
    ("期末", 3),
]


def _resolve(scheme_item, key):
    if key == "color":
        name = scheme_item.get("color", "黑色")
        return COLOR_REF.get(name, BLACK)
    if key == "size":
        name = scheme_item.get("size", "五号")
        return SIZE_REF.get(name, FONT_WUHAO)
    return scheme_item.get(key, None)


def _exam_sort_key(item):
    school, exam, content = item
    order = 99
    for kw, val in EXAM_ORDER_KEYWORDS:
        if kw in exam:
            order = val
            break
    return (order, school or "", exam or "")


def _exam_num_map(exam_display):
    if not exam_display:
        return "1"
    if "第一次月考" in exam_display:
        return "1"
    if "第二次月考" in exam_display:
        return "3"
    if "月考" in exam_display:
        return "1"
    if "期中" in exam_display:
        return "2"
    if "期末" in exam_display:
        return "4"
    return "1"


from .classify import classify_global, classify_gushici, classify_nonpoem


def _classify_text(text, type_name):
    result = classify_global(text)
    if result is not None:
        return result
    if type_name == "古诗词阅读":
        return classify_gushici(text)
    if type_name in ("论述类文本", "文学类文本", "文言文阅读"):
        return classify_nonpoem(text)
    return "other"


class BaseFormatter:

    def __init__(self, doc, format_scheme, xdf_mode, type_name):
        self.doc = doc
        self.xdf_mode = xdf_mode
        self.type_name = type_name

        self.READING_FONT = format_scheme["reading"].get("font", "楷体")
        self.READING_SIZE = _resolve(format_scheme["reading"], "size")
        self.READING_COLOR = _resolve(format_scheme["reading"], "color")
        self.READING_INDENT = format_scheme["reading"].get("indent", True)

        self.QUESTION_FONT = format_scheme["question"].get("font", "宋体")
        self.QUESTION_SIZE = _resolve(format_scheme["question"], "size")
        self.QUESTION_COLOR = _resolve(format_scheme["question"], "color")

        self.ANSWER_FONT = format_scheme["answer"].get("font", "宋体")
        self.ANSWER_SIZE = _resolve(format_scheme["answer"], "size")
        self.ANSWER_COLOR = _resolve(format_scheme["answer"], "color")
        self.ANSWER_BOLD = format_scheme["answer"].get("bold", False)

        self.EXPLAIN_FONT = self.ANSWER_FONT
        self.EXPLAIN_SIZE = self.ANSWER_SIZE
        self.EXPLAIN_COLOR = self.ANSWER_COLOR
        self.EXPLAIN_BOLD = self.ANSWER_BOLD

        self.ZONE_SCHEMES = {}

        _reading_defaults = {
            "font": self.QUESTION_FONT, "size": self.QUESTION_SIZE,
            "color": self.QUESTION_COLOR, "bold": False,
            "align": "left", "indent": False, "skip": False,
        }
        _reading_body_defaults = {
            "font": self.READING_FONT, "size": self.READING_SIZE,
            "color": self.READING_COLOR, "bold": False,
            "align": "left", "indent": self.READING_INDENT, "skip": False,
        }
        _title_defaults = {
            "font": self.READING_FONT, "size": self.READING_SIZE,
            "color": self.READING_COLOR, "bold": True,
            "align": "center", "indent": False, "skip": False,
        }
        _center_defaults = {
            "font": self.READING_FONT, "size": self.READING_SIZE,
            "color": self.READING_COLOR, "bold": False,
            "align": "center", "indent": False, "skip": False,
        }
        for zn, defaults in [
            ("reading_instruction", _reading_defaults),
            ("reading_body", _reading_body_defaults),
            ("book_title", _title_defaults),
            ("poem_text", _center_defaults),
            ("author", _center_defaults),
            ("material_marker", _reading_defaults),
            ("source_attribution", _reading_defaults),
            ("annotation", _reading_defaults),
            ("subsection_title", _title_defaults),
            ("source_label", _reading_defaults),
        ]:
            z = format_scheme.get(zn, {})
            self.ZONE_SCHEMES[zn] = {
                "font": z.get("font", defaults["font"]),
                "size": _resolve(z, "size") if "size" in z else defaults["size"],
                "color": _resolve(z, "color") if "color" in z else defaults["color"],
                "bold": z.get("bold", defaults["bold"]),
                "align": z.get("align", defaults["align"]),
                "indent": z.get("indent", defaults["indent"]),
                "skip": z.get("skip", defaults["skip"]),
            }

        _q_defaults = {
            "font": self.QUESTION_FONT, "size": self.QUESTION_SIZE,
            "color": self.QUESTION_COLOR, "bold": False,
            "align": "left", "indent": False, "skip": False,
        }
        for zn in ["option", "sub_question"]:
            z = format_scheme.get(zn, {})
            self.ZONE_SCHEMES[zn] = {
                "font": z.get("font", _q_defaults["font"]),
                "size": _resolve(z, "size") if "size" in z else _q_defaults["size"],
                "color": _resolve(z, "color") if "color" in z else _q_defaults["color"],
                "bold": z.get("bold", _q_defaults["bold"]),
                "align": z.get("align", _q_defaults["align"]),
                "indent": z.get("indent", _q_defaults["indent"]),
                "skip": z.get("skip", _q_defaults["skip"]),
            }

    def _apply_zone_format(self, para, text, zone_name, runs_data=None):
        z = self.ZONE_SCHEMES.get(zone_name)
        if z is None:
            return False
        if z["skip"]:
            return True
        if runs_data:
            for rt, _, ru in runs_data:
                cleaned = rt.replace('\u3000', '')
                add_run_with_font(para, cleaned, font_name=z["font"],
                                  font_size=z["size"],
                                  bold=z["bold"],
                                  color=z["color"],
                                  underline=ru)
        else:
            add_run_with_font(para, text, font_name=z["font"],
                              font_size=z["size"],
                              bold=z["bold"],
                              color=z["color"])
        if z["align"] == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif z["align"] == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if z["indent"]:
            para.paragraph_format.first_line_indent = Cm(0.74)
        return True

    ZONE_HANDLER_MAP = {
        "reading": "_format_reading",
        "question": "_format_question",
        "answer": "_format_answer",
        "explanation": "_format_explanation",
        "source_label": "_format_question",
        "skip": None,
    }

    def format_all(self, zones, content):
        for zone_name, zs, ze in zones:
            if zs > ze:
                continue
            poetry_idx = 0
            for li in range(zs, ze + 1):
                raw = content[li]
                text = raw[0]
                for ch in ['\u2b07', '\u000b', '\u000c', '\u0085']:
                    text = text.replace(ch, '\n')
                text = re.sub(r'[\ufffc\ufffd\ufeff\u200b-\u200d\u2060-\u2064\u2610-\u2612\u25a0\u25a1\u2b1c\u2b1d\u2b55\u3000]', '', text)
                text = text.strip()
                is_bold = raw[1]
                font_size = raw[2]
                runs_data = raw[3] if len(raw) > 3 else None
                orig_align = raw[4] if len(raw) > 4 else None
                if not text:
                    self.doc.add_paragraph()
                    continue

                segments = text.split('\n') if '\n' in text else [text]
                seg_runs = runs_data if len(segments) == 1 else None

                _score_pat = re.compile(r'^\d+[．.、]?(?:[（(]\d+分[）)])?$')
                merged = []
                i = 0
                while i < len(segments):
                    seg = segments[i]
                    s = seg.strip()
                    if (s and _score_pat.match(s)
                        and i + 1 < len(segments)
                        and '\u9605\u8bfb' in segments[i + 1]):
                        merged.append(seg + segments[i + 1])
                        i += 2
                    else:
                        merged.append(seg)
                        i += 1
                segments = merged
                if len(segments) == 1 and runs_data:
                    seg_runs = runs_data

                for seg in segments:
                    seg = seg.strip()
                    if not seg:
                        self.doc.add_paragraph()
                        continue

                    cls_global = _classify_text(seg, self.type_name)
                    if cls_global in ("skip", "skip_header", "skip_content") and not self.xdf_mode:
                        continue
                    if self.xdf_mode and cls_global in ("skip", "skip_header", "skip_content"):
                        para = self.doc.add_paragraph()
                        add_run_with_font(para, seg, font_name=self.QUESTION_FONT,
                                          font_size=self.QUESTION_SIZE,
                                          color=self.QUESTION_COLOR, bold=True)
                        continue

                    handled = False
                    handler_name = self.ZONE_HANDLER_MAP.get(zone_name)
                    if cls_global in ("answer_marker",):
                        handler_name = self.ZONE_HANDLER_MAP.get("answer")
                    elif cls_global in ("explanation_marker", "explanation_auto"):
                        handler_name = self.ZONE_HANDLER_MAP.get("explanation")
                    effective_zone = zone_name
                    if (self.type_name == "古诗词阅读" and
                        handler_name == "_format_reading" and
                        cls_global in ("poem_text", "poem_title", "poem_author", "other")):
                        if poetry_idx == 0:
                            effective_zone = "book_title"
                        elif poetry_idx == 1:
                            effective_zone = "author"
                        poetry_idx += 1

                    if handler_name == "_format_reading":
                        handled = self._format_reading(seg, cls_global, effective_zone, seg_runs, orig_align)
                    elif handler_name == "_format_question":
                        handled = self._format_question(seg, cls_global, zone_name, seg_runs, orig_align)
                    elif handler_name == "_format_answer":
                        handled = self._format_answer(seg, seg_runs, orig_align)
                    elif handler_name == "_format_explanation":
                        handled = self._format_explanation(seg, seg_runs, orig_align)
                    elif handler_name is None:
                        continue
                    if not handled:
                        para = self.doc.add_paragraph()
                        if self.READING_INDENT:
                            para.paragraph_format.first_line_indent = Cm(0.74)
                        if seg_runs:
                            for rt, _, ru in seg_runs:
                                add_run_with_font(para, rt.replace('\u3000', ''),
                                                  font_name=self.READING_FONT,
                                                  font_size=self.READING_SIZE,
                                                  color=self.READING_COLOR, bold=False,
                                                  underline=ru)
                        else:
                            add_run_with_font(para, seg, font_name=self.READING_FONT,
                                              font_size=self.READING_SIZE,
                                              color=self.READING_COLOR)

    @staticmethod
    def _clean_lines(text):
        if '\n' in text:
            return '\n'.join(line.strip() for line in text.split('\n'))
        return text

    def _format_reading(self, text, cls_r, zone_name="reading", runs_data=None, orig_align=None):
        stripped = text.strip()



        if zone_name in self.ZONE_SCHEMES:
            scheme = self.ZONE_SCHEMES[zone_name]
            para = self.doc.add_paragraph()
            if self._apply_zone_format(para, text, zone_name, runs_data):
                return True

        if re.match(r'^[（(][一二三四五六七八九十百]+[）)]', stripped):
            if self.xdf_mode:
                para = self.doc.add_paragraph()
                add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR, bold=True)
                return True
            else:
                return True

        m = re.match(r'^\d+[．.、]', stripped)
        has_number = bool(m)
        after_num = stripped[m.end():].lstrip() if m else ""
        is_inst = "阅读" in after_num
        is_score = bool(re.search(r'[（(]\d+分[）)]', after_num))
        if has_number and (cls_r == "instruction" or
            (cls_r in ("question", "modern_text") and (is_inst or is_score) and len(text.split('\n')) <= 2)):
            sub_lines = text.split('\n')
            if len(sub_lines) >= 2:
                last_line = sub_lines[-1].strip()
                if last_line and len(last_line) <= 30 and \
                   not re.search(r'[。！？；，、：]', last_line) and \
                   not re.match(r'^[\d（(【]', last_line):
                    inst_text = '\n'.join(sub_lines[:-1]).strip()
                    if inst_text:
                        para = self.doc.add_paragraph()
                        add_run_with_font(para, inst_text,
                                          font_name=self.QUESTION_FONT,
                                          font_size=self.QUESTION_SIZE,
                                          color=self.QUESTION_COLOR)
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run_with_font(para, last_line,
                                      font_name=self.READING_FONT,
                                      font_size=self.READING_SIZE,
                                      color=self.READING_COLOR, bold=True)
                    return True
            para = self.doc.add_paragraph()
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if cls_r == "instruction":
            para = self.doc.add_paragraph()
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if cls_r == "annotation" or stripped.startswith('【注】') or \
           stripped.startswith('[注]') or \
           stripped.startswith('注：') or stripped.startswith('注:'):
            para = self.doc.add_paragraph()
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if cls_r == "source_label":
            if self.xdf_mode:
                para = self.doc.add_paragraph()
                add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR, bold=True)
                return True
            else:
                return True

        if re.match(r'^[（(]\d+[）)]', stripped):
            para = self.doc.add_paragraph()
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if re.match(r'^[（(]?(摘编自|节选自|选自|摘自|有删改|有删节)', stripped) or \
           '有删改' in stripped or '有删节' in stripped:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_with_font(para, stripped, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if stripped.startswith('——') or stripped.startswith('—\u300a'):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_with_font(para, stripped, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if re.match(r'^(材料|文本|选文)[一二三四五六七八九十\d][：:]?\s*$', stripped):
            para = self.doc.add_paragraph()
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        title_max_len = 40
        if len(stripped) <= title_max_len and not re.search(r'[。！？；]', stripped):
            is_book_title = (
                (stripped.startswith('《') and '》' in stripped)
                or ('节选' in stripped)
                or (len(stripped) <= 20 and
                    orig_align == WD_ALIGN_PARAGRAPH.CENTER and
                    not re.search(r'[\d+．.、（）()【】""''「」『』…—]', stripped))
            )
            if is_book_title:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_with_font(para, text, font_name=self.READING_FONT,
                                  font_size=self.READING_SIZE,
                                  color=self.READING_COLOR,
                                  bold=True)
                return True

        if self.xdf_mode:
            if re.match(r'^材料[一二三四五六七八九十\d]', stripped):
                para = self.doc.add_paragraph()
                add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR)
                return True
            if stripped.startswith('（摘编自') or stripped.startswith('摘编自') or \
               stripped.startswith('（摘自') or stripped.startswith('摘自') or \
               stripped.startswith('（有删改') or stripped.startswith('有删改') or \
               stripped.startswith('（有删节') or stripped.startswith('有删节') or \
               '有删改' in stripped or '有删节' in stripped:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR)
                return True
            if stripped.startswith('——') or stripped.startswith('—\u300a'):
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR)
                return True

        if cls_r == "modern_text" and has_number and is_inst:
            sub_lines = text.split('\n')
            in_reading = False
            for sub_line in sub_lines:
                sl = sub_line.strip()
                if not sl:
                    continue
                if not in_reading:
                    if re.match(r'^\d+[．.、]', sl):
                        para = self.doc.add_paragraph()
                        add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                          font_size=self.QUESTION_SIZE,
                                          color=self.QUESTION_COLOR)
                        continue
                    m2 = re.match(r'^(\d+[．.、])', sl)
                    if m2 and sl[m2.end():].lstrip().startswith("阅读"):
                        para = self.doc.add_paragraph()
                        add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                          font_size=self.QUESTION_SIZE,
                                          color=self.QUESTION_COLOR)
                        in_reading = True
                        continue
                in_reading = True
                if re.match(r'^[（(]?(摘编自|节选自|选自|摘自|有删改|有删节)', sl) or \
                   '有删改' in sl or '有删节' in sl:
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                      font_size=self.QUESTION_SIZE,
                                      color=self.QUESTION_COLOR)
                    continue
                if sl.startswith('——') or sl.startswith('—\u300a'):
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                      font_size=self.QUESTION_SIZE,
                                      color=self.QUESTION_COLOR)
                    continue
                if re.match(r'^(?:[［][\u4e00-\u9fff]{1,3}[］]|\[[\u4e00-\u9fff]{1,3}\])', sl) or \
                   sl.startswith('[注]') or sl.startswith('【注】') or \
                   sl.startswith('[注释]') or sl.startswith('【注释】'):
                    para = self.doc.add_paragraph()
                    add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                      font_size=self.QUESTION_SIZE,
                                      color=self.QUESTION_COLOR)
                    continue
                para = self.doc.add_paragraph()
                if self.READING_INDENT:
                    para.paragraph_format.first_line_indent = Cm(0.74)
                add_run_with_font(para, sl, font_name=self.READING_FONT,
                                  font_size=self.READING_SIZE,
                                  color=self.READING_COLOR)
            return True

        if '\n' in stripped:
            for sep in ['【注】', '【注释】', '[注]', '[注释]']:
                if sep in stripped:
                    prefix = stripped.split(sep)[0].strip()
                    clean_prefix = re.sub(r'[\s，。！？、；：\u201c\u201d\u2018\u2019《》（）\u2460-\u2473]', '', prefix)
                    is_poem = 4 <= len(clean_prefix) <= 28
                    if not is_poem:
                        idx = stripped.index(sep)
                        before = stripped[:idx].strip()
                        after = stripped[idx:].strip()
                        if before:
                            para = self.doc.add_paragraph()
                            if self.READING_INDENT:
                                para.paragraph_format.first_line_indent = Cm(0.74)
                            add_run_with_font(para, before, font_name=self.READING_FONT,
                                              font_size=self.READING_SIZE, color=self.READING_COLOR)
                        para = self.doc.add_paragraph()
                        add_run_with_font(para, after, font_name=self.QUESTION_FONT,
                                          font_size=self.QUESTION_SIZE, color=self.QUESTION_COLOR)
                        return True
                    break

        if re.match(r'^(?:[［][\u4e00-\u9fff]{1,3}[］]|\[[\u4e00-\u9fff]{1,3}\])', stripped) or \
           stripped.startswith('[注]') or stripped.startswith('【注】') or \
           stripped.startswith('[注释]') or stripped.startswith('【注释】'):
            para = self.doc.add_paragraph()
            add_run_with_font(para, stripped, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
            return True

        if self.type_name == "古诗词阅读":
            stripped_poem = stripped
            sub_parts = re.split(r'[\n\v\f\r\u2028\u2029\u0085]', stripped_poem)
            has_subs = len([s for s in sub_parts if s.strip()]) > 1

            if has_subs:
                poem_line = sub_parts[0].strip()
                poem_clean = re.sub(r'[\u00b7\u2022\u25cf\u2b07\ufffc\ufffd]', '', poem_line)
                clean_first = re.sub(r'[\s，。！？、；：\u201c\u201d\u2018\u2019《》（）\u2460-\u2473]', '', poem_clean)
                if 4 <= len(clean_first) <= 28 and len(clean_first) >= len(poem_clean) * 0.5:
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run_with_font(para, poem_clean, font_name=self.READING_FONT,
                                      font_size=self.READING_SIZE, color=self.READING_COLOR)
                    for sub_line in sub_parts[1:]:
                        sl = sub_line.strip()
                        if not sl or len(sl) <= 2:
                            continue
                        if sl.startswith('[注]') or sl.startswith('【注】') or \
                           sl.startswith('[注释]') or sl.startswith('【注释】') or \
                           re.match(r'^(?:[［][\u4e00-\u9fff]{1,3}[］]|\[[\u4e00-\u9fff]{1,3}\])', sl):
                            para = self.doc.add_paragraph()
                            add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                              font_size=self.QUESTION_SIZE, color=self.QUESTION_COLOR)
                        elif re.match(r'^[（(]?(摘编自|节选自|选自|摘自|有删改|有删节)', sl) or \
                             '有删改' in sl or '有删节' in sl:
                            para = self.doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                              font_size=self.QUESTION_SIZE, color=self.QUESTION_COLOR)
                        elif sl.startswith('——') or sl.startswith('—\u300a'):
                            para = self.doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            add_run_with_font(para, sl, font_name=self.QUESTION_FONT,
                                              font_size=self.QUESTION_SIZE, color=self.QUESTION_COLOR)
                        else:
                            para = self.doc.add_paragraph()
                            if self.READING_INDENT:
                                para.paragraph_format.first_line_indent = Cm(0.74)
                            add_run_with_font(para, sl, font_name=self.READING_FONT,
                                              font_size=self.READING_SIZE, color=self.READING_COLOR)
                    return True

            clean = re.sub(r'[\s，。！？、；：\u201c\u201d\u2018\u2019《》（）]', '', stripped_poem)
            if 4 <= len(clean) <= 28 and len(clean) >= len(stripped_poem) * 0.5:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_with_font(para, stripped_poem, font_name=self.READING_FONT,
                                  font_size=self.READING_SIZE,
                                  color=self.READING_COLOR)
                return True
            if re.match(r'^[\u4e00-\u9fff·\s]{2,10}$', stripped_poem):
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_with_font(para, stripped_poem, font_name=self.READING_FONT,
                                  font_size=self.READING_SIZE,
                                  color=self.READING_COLOR,
                                  bold=True)
                return True

        if orig_align == WD_ALIGN_PARAGRAPH.CENTER and len(stripped) <= 40 and \
           not re.search(r'[。！？；\u2014]', stripped):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_font(para, text, font_name=self.READING_FONT,
                              font_size=self.READING_SIZE,
                              color=self.READING_COLOR, bold=True)
            return True

        if cls_r == "modern_text":
            para = self.doc.add_paragraph()
            if self.READING_INDENT:
                para.paragraph_format.first_line_indent = Cm(0.74)
            add_run_with_font(para, text, font_name=self.READING_FONT,
                              font_size=self.READING_SIZE,
                              color=self.READING_COLOR)
            return True

        return False

    def _format_question(self, text, cls_r, zone_name="question", runs_data=None, orig_align=None):

        if zone_name in self.ZONE_SCHEMES:
            para = self.doc.add_paragraph()
            if self._apply_zone_format(para, text, zone_name, runs_data):
                return True

        if cls_r == "instruction":
            return True
        if cls_r in ("skip", "skip_header", "skip_content"):
            return True

        if '本题考查' in text or text.startswith('【分析】') or \
           re.match(r'^故选', text):
            para = self.doc.add_paragraph()
            if runs_data:
                for rt, _, ru in runs_data:
                    add_run_with_font(para, rt, font_name=self.EXPLAIN_FONT,
                                      font_size=self.EXPLAIN_SIZE,
                                      color=self.EXPLAIN_COLOR, bold=self.EXPLAIN_BOLD,
                                      underline=ru)
            else:
                add_run_with_font(para, text, font_name=self.EXPLAIN_FONT,
                                  font_size=self.EXPLAIN_SIZE,
                                  color=self.EXPLAIN_COLOR)
            return True

        para = self.doc.add_paragraph()
        if runs_data:
            for rt, _, ru in runs_data:
                add_run_with_font(para, rt, font_name=self.QUESTION_FONT,
                                  font_size=self.QUESTION_SIZE,
                                  color=self.QUESTION_COLOR, bold=False,
                                  underline=ru)
        else:
            add_run_with_font(para, text, font_name=self.QUESTION_FONT,
                              font_size=self.QUESTION_SIZE,
                              color=self.QUESTION_COLOR)
        return True

    @staticmethod
    def _strip_prefix_from_runs(runs_data, prefix_len):
        remaining = []
        consumed = 0
        for rt, rb, ru in runs_data:
            if consumed < prefix_len:
                to_consume = min(len(rt), prefix_len - consumed)
                rest = rt[to_consume:]
                consumed += to_consume
                if rest:
                    remaining.append((rest, rb, ru))
            else:
                remaining.append((rt, rb, ru))
        return remaining

    def _format_answer(self, text, runs_data=None, orig_align=None):
        text = self._clean_lines(text)
        segments = text.split('\n')
        for i, seg in enumerate(segments):
            seg = seg.strip()
            if not seg:
                self.doc.add_paragraph()
                continue
            seg_runs = runs_data if len(segments) == 1 else None
            para = self.doc.add_paragraph()
            bold = self.ANSWER_BOLD
            marker = "【答案】"
            if seg.startswith(marker):
                add_run_with_font(para, marker, font_name=self.ANSWER_FONT,
                                  font_size=self.ANSWER_SIZE,
                                  bold=bold, color=self.ANSWER_COLOR)
                remain = seg[len(marker):].strip()
                if remain:
                    if seg_runs:
                        remaining_runs = self._strip_prefix_from_runs(seg_runs, len(marker))
                        for rt, _, ru in remaining_runs:
                            add_run_with_font(para, rt.replace('\u3000', ''), font_name=self.ANSWER_FONT,
                                              font_size=self.ANSWER_SIZE,
                                              color=self.ANSWER_COLOR, bold=bold,
                                              underline=ru)
                    else:
                        add_run_with_font(para, remain, font_name=self.ANSWER_FONT,
                                          font_size=self.ANSWER_SIZE,
                                          color=self.ANSWER_COLOR, bold=bold)
                continue
            if seg_runs:
                for rt, _, ru in seg_runs:
                    add_run_with_font(para, rt.replace('\u3000', ''), font_name=self.ANSWER_FONT,
                                      font_size=self.ANSWER_SIZE,
                                      color=self.ANSWER_COLOR, bold=bold,
                                      underline=ru)
            else:
                add_run_with_font(para, seg, font_name=self.ANSWER_FONT,
                                  font_size=self.ANSWER_SIZE,
                                  color=self.ANSWER_COLOR, bold=bold)
        return True

    def _format_explanation(self, text, runs_data=None, orig_align=None):
        text = self._clean_lines(text)
        segments = text.split('\n')
        for i, seg in enumerate(segments):
            seg = seg.strip()
            if not seg:
                self.doc.add_paragraph()
                continue
            seg_runs = runs_data if len(segments) == 1 else None
            para = self.doc.add_paragraph()
            bold = self.EXPLAIN_BOLD
            for marker in ["【解析】", "【导语】", "【详解】", "【点睛】"]:
                if seg.startswith(marker):
                    add_run_with_font(para, marker, font_name=self.EXPLAIN_FONT,
                                      font_size=self.EXPLAIN_SIZE,
                                      bold=bold, color=self.EXPLAIN_COLOR)
                    remain = seg[len(marker):].strip()
                    if remain:
                        if seg_runs:
                            remaining_runs = self._strip_prefix_from_runs(seg_runs, len(marker))
                            for rt, _, ru in remaining_runs:
                                add_run_with_font(para, rt.replace('\u3000', ''), font_name=self.EXPLAIN_FONT,
                                                  font_size=self.EXPLAIN_SIZE,
                                                  color=self.EXPLAIN_COLOR, bold=bold,
                                                  underline=ru)
                        else:
                            add_run_with_font(para, remain, font_name=self.EXPLAIN_FONT,
                                              font_size=self.EXPLAIN_SIZE,
                                              color=self.EXPLAIN_COLOR, bold=bold)
                    break
            else:
                if seg_runs:
                    for rt, _, ru in seg_runs:
                        add_run_with_font(para, rt.replace('\u3000', ''), font_name=self.EXPLAIN_FONT,
                                          font_size=self.EXPLAIN_SIZE,
                                          color=self.EXPLAIN_COLOR, bold=bold,
                                          underline=ru)
                else:
                    add_run_with_font(para, seg, font_name=self.EXPLAIN_FONT,
                                      font_size=self.EXPLAIN_SIZE,
                                      color=self.EXPLAIN_COLOR, bold=bold)
        return True
