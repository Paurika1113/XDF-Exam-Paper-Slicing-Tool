"""
slice_all.py — 按4种题型分类输出DOCX
"""

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

from formatters import EXAM_ORDER_KEYWORDS, FORMATTER_MAP, _classify_text
from formatters.classify import classify_global, classify_gushici, classify_nonpoem
from scanner import scan_with_fallback


def get_paper_ranges(doc, fname):
    result = scan_with_fallback(doc, fname)
    if result and len(result) >= 3:
        return result

    print("  [!] \u901a\u7528\u68c0\u6d4b\u672a\u5b8c\u5168\u8986\u76d6\uff0c\u56de\u9000\u5230\u81ea\u52a8\u68c0\u6d4b")
    top_sections = []
    sub_sections = []

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        m1 = re.match(r'^[一二三四五六七八九十百]+[、.．]\s*(.*)', t)
        if m1:
            top_sections.append((i, m1.group(1).strip() or t[:40], t[:60]))
            continue
        m2 = re.match(r'^[（(][一二三四五六七八九十百]+[）)]\s*(.*)', t)
        if m2:
            sub_sections.append((i, m2.group(1).strip() or t[:40], t[:60]))
            continue

    total = len(doc.paragraphs)

    def next_section_idx(cur):
        candidates = []
        for pi in range(cur + 1, total):
            t = doc.paragraphs[pi].text.strip()
            if not t:
                continue
            if re.match(r'^【', t) or re.match(r'^\[', t):
                candidates.append(pi)
        return min(candidates) if candidates else total

    def _safe_end(start):
        end = next_section_idx(start)
        if end > start + 1:
            prev = doc.paragraphs[end - 1].text.strip()
            if any(kw in prev for kw in ["阅读下面的文字", "阅读下面文字", "阅读下列文字"]):
                end -= 1
        return end

    ranges = {}

    for idx, name, full in (sub_sections + top_sections):
        if any(kw in full for kw in ['阅读Ⅰ', '阅读I', '阅读1', '现代文阅读Ⅰ', '现代文阅读I']):
            end = _safe_end(idx)
            ranges["论述类文本"] = (idx, end - 1)
            break
    if "论述类文本" not in ranges:
        for idx, name, full in top_sections:
            if re.match(r'^[一二三四五六七八九十百]+[、.．]', full) and '阅读' in full:
                top_idx = idx
                sub_within = [(si, sf) for si, sn, sf in sub_sections if si > top_idx and si < next_section_idx(top_idx)]
                if sub_within:
                    first_sub = sub_within[0]
                    if 'Ⅱ' not in first_sub[2] and 'II' not in first_sub[2]:
                        end = _safe_end(first_sub[0])
                        ranges["论述类文本"] = (first_sub[0], end - 1)
                break

    for idx, name, full in (sub_sections + top_sections):
        if any(kw in full for kw in ['阅读Ⅱ', '阅读II', '阅读2', '现代文阅读Ⅱ', '现代文阅读II']):
            end = _safe_end(idx)
            ranges["文学类文本"] = (idx, end - 1)
            break
    if "文学类文本" not in ranges and "论述类文本" in ranges:
        ranges["文学类文本"] = ranges["论述类文本"]
        print("  [!] 文学类文本复用论述类范围")

    for idx, name, full in (sub_sections + top_sections):
        if any(kw in full for kw in ['文言文', '文言知识']):
            end = _safe_end(idx)
            ranges["文言文阅读"] = (idx, end - 1)
            break
    if "文言文阅读" not in ranges and '同安' in fname:
        for idx, name, full in (sub_sections + top_sections):
            if '阅读Ⅲ' in full or '阅读III' in full:
                end = _safe_end(idx)
                ranges["文言文阅读"] = (idx, end - 1)
                break

    for idx, name, full in (sub_sections + top_sections):
        if any(kw in full for kw in ['古代诗歌阅读', '古代诗歌鉴赏', '诗歌阅读', '诗歌鉴赏', '阅读Ⅳ', '阅读IV']):
            if '古诗文阅读' in full:
                continue
            end = _safe_end(idx)
            ranges["古诗词阅读"] = (idx, end - 1)
            break
    if "古诗词阅读" not in ranges:
        for idx, name, full in top_sections:
            if '古代诗歌' in full or '古诗' in full:
                end = _safe_end(idx)
                ranges["古诗词阅读"] = (idx, end - 1)
                break
    if "古诗词阅读" not in ranges:
        for idx, name, full in sub_sections:
            if '阅读Ⅳ' in full or '阅读IV' in full:
                end = _safe_end(idx)
                ranges["古诗词阅读"] = (idx, end - 1)
                break

    return ranges


def extract_type_content(doc, start_idx, end_idx):
    blocks = []
    for i in range(start_idx, min(end_idx + 1, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text
        text = text.replace('\u2b07', '\n')
        text = text.replace('\u000b', '\n')
        text = text.replace('\u000c', '\n')
        text = text.replace('\u0085', '\n')
        if not text.strip():
            blocks.append(("", None, None, None, None))
            continue

        any_bold = None
        font_sizes = []
        runs_data = []
        for run in p.runs:
            rt = run.text
            if not rt:
                continue
            rb = run.bold
            ru = run.font.underline
            runs_data.append((rt, rb, ru))
            if rb:
                any_bold = True
            elif any_bold is None:
                any_bold = False
            if run.font.size:
                font_sizes.append(run.font.size)
        font_size = max(font_sizes) if font_sizes else None

        blocks.append((text.strip(), any_bold, font_size, runs_data if runs_data else None, p.alignment))

    return blocks


def _add_run_with_font(paragraph, text, font_name='宋体', font_size=None,
                        bold=False, color=None, italic=False, underline=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if underline is not None:
        run.font.underline = underline
    if color:
        run.font.color.rgb = color
    return run


def _add_shading(paragraph, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    paragraph._element.get_or_add_pPr().append(shading)


def detect_interleaved(classified, lines):
    """检测交替格式：answer/explanation 区之后是否出现有【答案】跟随的（N）题目。"""
    for i, cls in enumerate(classified):
        if cls in ("answer_marker", "explanation_marker", "explanation_auto"):
            for j in range(i + 1, len(classified)):
                t = lines[j][0].strip()
                if re.match(r'^[（(]\d+[）)]', t) and classified[j] == "question":
                    for k in range(j + 1, len(classified)):
                        nt = lines[k][0].strip()
                        if re.match(r'^[（(]\d+[）)]', nt):
                            break
                        if classified[k] == "answer_marker":
                            return True
                    break
            break
    return False


def find_zones(lines, type_name=None, interleaved=False):
    total = len(lines)
    zones = []
    current = "reading"
    z_start = 0
    exam_type_pat = re.compile(r'^[（(]\d*[）)]\s*(期中|期末|月考)')

    classified = []
    for text, _, _, *_ in lines:
        classified.append(_classify_text(text, type_name or "文言文阅读"))

    for i, (text, _, _, *_) in enumerate(lines):
        t = text.strip()
        if not t:
            continue
        cls = classified[i]

        if current == "source_label":
            zones.append(("source_label", z_start, i - 1))
            current = "reading"
            z_start = i

        if current == "question" and \
            (cls == "skip" or cls == "instruction" or (cls == "other" and (
                re.match(r'^[（(][一二三四五六七八九十百]+[）)]', t) or
                re.match(r'^【', t) or
                re.match(r'^[（(]\d+[）)]', t)))):
            zones.append(("question", z_start, i - 1))
            current = "reading"
            z_start = i
            continue

        if current == "reading":
            if re.match(r'^\d+[．.、]', t):
                m = re.match(r'^(\d+[．.、])', t)
                after_num = t[m.end():].lstrip() if m else ""
                is_inst = "阅读" in after_num
                is_score = bool(re.search(r'[（(]\d+分[）)]', after_num))
                if '本题考查' in t:
                    zones.append(("reading", z_start, i - 1))
                    zones.append(("explanation", i, i))
                    z_start = i + 1
                    current = "explanation"
                elif is_inst or is_score:
                    pass
                else:
                    zones.append(("reading", z_start, i - 1))
                    current = "question"
                    z_start = i
                continue

            if cls == "answer_marker":
                zones.append(("reading", z_start, i - 1))
                current = "answer"
                z_start = i
                continue

            if cls in ("explanation_marker", "explanation_auto"):
                zones.append(("reading", z_start, i - 1))
                current = "explanation"
                z_start = i
                continue

            if cls == "question" and re.match(r'^[（(]\d+[）)]', t) and not \
               any(kw in t for kw in ['月考', '期中', '期末', '考试']):
                zones.append(("reading", z_start, i - 1))
                current = "question"
                z_start = i
                continue

            if cls == "source_label" or exam_type_pat.match(t):
                zones.append(("reading", z_start, i - 1))
                current = "source_label"
                z_start = i
                continue

            continue

        if current == "question":
            if cls == "source_label" or exam_type_pat.match(t):
                zones.append(("question", z_start, i - 1))
                current = "source_label"
                z_start = i
                continue
            if t.startswith("【答案】"):
                zones.append(("question", z_start, i - 1))
                current = "answer"
                z_start = i
                continue
            if cls in ("explanation_marker", "explanation_auto") or '本题考查' in t:
                zones.append(("question", z_start, i - 1))
                current = "explanation"
                z_start = i
                continue
            continue

        if current == "answer":
            if cls == "source_label" or exam_type_pat.match(t):
                zones.append(("answer", z_start, i - 1))
                current = "source_label"
                z_start = i
                continue
            if cls in ("explanation_marker", "explanation_auto"):
                zones.append(("answer", z_start, i - 1))
                current = "explanation"
                z_start = i
                continue
            if interleaved and cls == "question" and re.match(r'^[（(]\d+[）)]', t):
                zones.append(("answer", z_start, i - 1))
                current = "question"
                z_start = i
                continue
            continue

        if current == "explanation":
            if cls == "source_label" or exam_type_pat.match(t):
                zones.append(("explanation", z_start, i - 1))
                current = "source_label"
                z_start = i
                continue
            if re.match(r'^\d+[．.、]', t):
                if '本题考查' in t or cls in ("explanation_auto",) or \
                   re.match(r'^\d+[．.、]\S+：', t) or \
                   re.search(r'[。，]\s*\d+[．.、]', t):
                    pass
                else:
                    zones.append(("explanation", z_start, i - 1))
                    current = "question"
                    z_start = i
                continue
            if interleaved and cls == "question" and re.match(r'^[（(]\d+[）)]', t):
                zones.append(("explanation", z_start, i - 1))
                current = "question"
                z_start = i
                continue
            continue

    if z_start <= total - 1:
        zones.append((current, z_start, total - 1))

    return zones


def refine_zones(zones, content, type_name, interleaved=False):
    if interleaved:
        return zones
    first_answer_idx = None
    for i, (zt, _, _) in enumerate(zones):
        if zt == "answer":
            first_answer_idx = i
            break

    if first_answer_idx is not None and first_answer_idx < len(zones) - 1:
        for zt, _, _ in zones[first_answer_idx + 1:]:
            if zt == "reading":
                return zones
        before = zones[:first_answer_idx]
        remaining = zones[first_answer_idx:]
        merged = ("answer", remaining[0][1], remaining[-1][2])
        return before + [merged]

    return zones


PLATE_NAMES = {
    "论述类文本": "论述类阅读",
    "文学类文本": "文学类阅读",
    "文言文阅读": "文言文阅读",
    "古诗词阅读": "诗歌阅读",
}


def _build_exam_order_map(exam_order):
    if exam_order:
        return {v: i for i, v in enumerate(exam_order)}
    return {}


def _exam_sort_key(item, exam_order_map=None):
    school, exam = item[0], item[1]
    if exam_order_map:
        i = exam_order_map.get(exam, -1)
        if i >= 0:
            return (i, school or "", exam or "")
    order = 99
    for kw, val in EXAM_ORDER_KEYWORDS:
        if kw in exam:
            order = val
            break
    return (order, school or "", exam or "")


def _add_row_tblPrEx(row):
    """给行添加 tblPrEx（边框+边距）+ trPr（居中）。"""
    row._tr.insert(0, parse_xml(
        f'<w:tblPrEx {nsdecls("w")}>'
        '  <w:tblBorders>'
        '    <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '    <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '    <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  </w:tblBorders>'
        '  <w:tblCellMar>'
        '    <w:top w:w="0" w:type="dxa"/>'
        '    <w:left w:w="108" w:type="dxa"/>'
        '    <w:bottom w:w="0" w:type="dxa"/>'
        '    <w:right w:w="108" w:type="dxa"/>'
        '  </w:tblCellMar>'
        '</w:tblPrEx>'))
    row._tr.insert(1, parse_xml(f'<w:trPr {nsdecls("w")}><w:jc w:val="center"/></w:trPr>'))


def _set_column_widths_from_content(table):
    """根据表格各列的最大内容宽度设置 gridCol 宽度。"""
    from docx.oxml.ns import qn
    avail_cm = 14.66  # A4 纵向 21cm - 3.17cm*2 页边距
    avail_twips = int(avail_cm / 2.54 * 1440)

    content_lens = [0, 0, 0, 0]
    for row in table.rows:
        for ci in range(4):
            text = row.cells[ci].text
            # 中文算 2 单位，拉丁算 1 单位
            w = sum(2 if ord(c) > 127 else 1 for c in text)
            content_lens[ci] = max(content_lens[ci], w)

    # 空列保底宽度
    min_lens = [6, 4, 2, 2]
    for ci in range(4):
        content_lens[ci] = max(content_lens[ci], min_lens[ci])

    total_units = sum(content_lens)
    if total_units == 0:
        return

    gc_list = list(table._tbl.tblGrid.findall(qn('w:gridCol')))
    widths = [int(avail_twips * cl / total_units) for cl in content_lens]
    diff = avail_twips - sum(widths)
    widths[0] += diff
    for ci, gc in enumerate(gc_list):
        gc.set(qn('w:w'), str(widths[ci]))
    for row in table.rows:
        for ci in range(4):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.get_or_add_tcW()
            tcW.w = widths[ci]
            tcW.type = 'dxa'


def generate_docx(type_name, items, output_path, grade_label="2025-2026高二上", format_scheme=None, exam_order=None):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    FONT_SANHAO = Pt(16)
    FONT_SIHAO = Pt(14)
    FONT_XIAOSI = Pt(12)
    FONT_WUHAO = Pt(10.5)
    RED = RGBColor(0xFF, 0x00, 0x00)
    BLUE = RGBColor(0x00, 0x00, 0xFF)
    BLACK = RGBColor(0x00, 0x00, 0x00)

    yr_match = re.search(r'(\d{4}-\d{4})', grade_label or '')
    table_year = yr_match.group(1) if yr_match else ""

    if not format_scheme or "reading" not in format_scheme:
        format_scheme = {
            "reading": {"font": "楷体", "size": FONT_WUHAO, "color": BLACK, "indent": True},
            "reading_body": {"font": "楷体", "size": FONT_WUHAO, "color": BLACK, "indent": True},
            "reading_instruction": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
             "book_title": {"size": FONT_WUHAO, "color": BLACK,
                           "align": "center"},
            "poem_text": {"font": "楷体", "size": FONT_WUHAO, "color": BLACK,
                          "align": "center"},
             "author": {"font": "楷体", "size": FONT_WUHAO, "color": BLACK,
                        "bold": True, "align": "center"},
            "material_marker": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "source_attribution": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK,
                                   "align": "right"},
            "annotation": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "subsection_title": {"skip": True},
            "source_label": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "question": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "option": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "sub_question": {"font": "宋体", "size": FONT_WUHAO, "color": BLACK},
            "answer": {"font": "宋体", "size": FONT_WUHAO, "color": "红色", "bold": False},
            "explanation": {"font": "宋体", "size": FONT_WUHAO, "color": "红色", "bold": False},
        }

    COLOR_REF = {"黑色": BLACK, "红色": RED, "蓝色": BLUE, "绿色": RGBColor(0x00,0x80,0x00)}
    SIZE_REF = {"小三": Pt(15), "四号": Pt(14), "小四": Pt(12), "五号": Pt(10.5), "小五": Pt(9)}

    def _resolve(scheme_item, key):
        if key == "color":
            name = scheme_item.get("color", "黑色")
            return COLOR_REF.get(name, BLACK)
        if key == "size":
            name = scheme_item.get("size", "五号")
            return SIZE_REF.get(name, FONT_WUHAO)
        return scheme_item.get(key, None)

    READING_FONT = format_scheme["reading"].get("font", "楷体")
    READING_SIZE = _resolve(format_scheme["reading"], "size")
    READING_COLOR = _resolve(format_scheme["reading"], "color")
    READING_INDENT = format_scheme["reading"].get("indent", True)

    QUESTION_FONT = format_scheme["question"].get("font", "宋体")
    QUESTION_SIZE = _resolve(format_scheme["question"], "size")
    QUESTION_COLOR = _resolve(format_scheme["question"], "color")

    ANSWER_FONT = format_scheme["answer"].get("font", "宋体")
    ANSWER_SIZE = _resolve(format_scheme["answer"], "size")
    ANSWER_COLOR = _resolve(format_scheme["answer"], "color")
    ANSWER_BOLD = format_scheme["answer"].get("bold", False)

    EXPLAIN_FONT = format_scheme["explanation"].get("font", "宋体")
    EXPLAIN_SIZE = _resolve(format_scheme["explanation"], "size")
    EXPLAIN_COLOR = _resolve(format_scheme["explanation"], "color")
    EXPLAIN_BOLD = format_scheme["explanation"].get("bold", False)

    xdf_mode = format_scheme.get("xdf_mode", False)

    plate = PLATE_NAMES.get(type_name, type_name)
    if xdf_mode:
        info_bar = doc.add_paragraph()
        info_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_with_font(info_bar, f"{grade_label}{plate}切片",
                           font_name='黑体', font_size=FONT_SANHAO, bold=True)
    else:
        info_bar = doc.add_paragraph()
        info_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_with_font(info_bar, f"{grade_label}{plate}切片",
                           font_name='黑体', font_size=FONT_SANHAO, bold=True)

    doc.add_paragraph()
    h1 = doc.add_paragraph()
    _add_run_with_font(h1, "一、考情一览", font_name='黑体', font_size=FONT_SIHAO, bold=True)

    _exam_order_map = _build_exam_order_map(exam_order)
    _sort_key = lambda it: _exam_sort_key(it, _exam_order_map)

    if exam_order:
        items.sort(key=_sort_key)
    school_list = sorted(set(it[0] for it in items))
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True

    tbl = table._tbl
    # 表格属性：100% 页宽 + 自动调整
    existing_tblW = tbl.tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tbl.tblPr.remove(existing_tblW)
    tbl.tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
    if tbl.tblPr.find(qn('w:autofit')) is None:
        tbl.tblPr.append(parse_xml(f'<w:autofit {nsdecls("w")}/>'))
    tbl.tblPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>')
    tbl.tblPr.append(tblBorders)
    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '  <w:top w:w="0" w:type="dxa"/>'
        '  <w:left w:w="108" w:type="dxa"/>'
        '  <w:bottom w:w="0" w:type="dxa"/>'
        '  <w:right w:w="108" w:type="dxa"/>'
        '</w:tblCellMar>')
    tbl.tblPr.append(tblCellMar)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '年度+学校'
    hdr_cells[1].text = '考试'
    hdr_cells[2].text = '文本类型'
    hdr_cells[3].text = '题型'

    _add_row_tblPrEx(table.rows[0])
    for cell in hdr_cells:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="CFCECE" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)

    sortable_items = []
    for it in items:
        school, exam, content = it[0], it[1], it[2]
        zones = it[3] if len(it) >= 4 else None
        title = it[4] if len(it) >= 5 else None
        sortable_items.append((school, exam, content, zones, title))
    sorted_items_raw = sorted(sortable_items, key=_sort_key)

    exam_groups = {}
    for item in sorted_items_raw:
        school, exam, content = item[0], item[1], item[2]
        for kw, _ in EXAM_ORDER_KEYWORDS:
            if kw in exam:
                exam_display = kw
                break
        else:
            exam_display = exam

        q_count = sum(1 for t, *_ in content
                      if re.match(r'^\d+[．.、]', t.strip()))

        key = (exam_display, school)
        if key not in exam_groups:
            exam_groups[key] = {"q_count": 0, "row": None}
        exam_groups[key]["q_count"] += max(q_count, 1)

    ordered_keys = sorted(exam_groups.keys(),
                          key=lambda k: (_sort_key((k[1], k[0], []))[0], k[1], k[0]))
    for exam_display, school in ordered_keys:
        new_row = table.add_row()
        row_cells = new_row.cells
        row_cells[0].text = f"{table_year}{school}" if table_year else school
        row_cells[1].text = exam_display
        row_cells[2].text = ''
        row_cells[3].text = ''
        _add_row_tblPrEx(new_row)
        for cell in row_cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 根据数据行的实际内容设置列宽
    _set_column_widths_from_content(table)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    _add_run_with_font(h2, "二、试题切片", font_name='黑体', font_size=FONT_SIHAO, bold=True)

    current_exam_display = None
    total_question_count = 0

    for idx, item in enumerate(sorted_items_raw):
        school, exam, content = item[0], item[1], item[2]
        zones = item[3] if len(item) >= 4 else None
        file_title = item[4] if len(item) >= 5 else None
        for kw, _ in EXAM_ORDER_KEYWORDS:
            if kw in exam:
                exam_display = kw
                break
        else:
            exam_display = exam

        if exam_display != current_exam_display:
            h3 = doc.add_paragraph()
            _add_run_with_font(h3, f"({_exam_num_map(exam_display)}){exam_display}",
                               font_name='黑体', font_size=FONT_XIAOSI, bold=True)
            current_exam_display = exam_display

        if zones is None:
            classified = [_classify_text(t, type_name or "文言文阅读")
                          for t, *_ in content]
            interleaved = detect_interleaved(classified, content)
            zones = find_zones(content, type_name, interleaved=interleaved)
            zones = refine_zones(zones, content, type_name, interleaved=interleaved)
        source_added = False

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        tag_text = f"【{file_title}】" if file_title else f"【{grade_label} {school} {exam_display}】"
        _add_run_with_font(para, tag_text,
                           font_name='宋体', font_size=FONT_WUHAO, bold=False)
        source_added = True

        formatter = FORMATTER_MAP[type_name](doc, format_scheme, xdf_mode, type_name)
        formatter.format_all(zones, content)

    doc.save(output_path)
    print(f"  [OK] {type_name}: {output_path.split(os.sep)[-1]} "
          f"({len(items)}\u4efd\u8bd5\u5377, {os.path.getsize(output_path) / 1024:.1f} KB)")


def _exam_num_map(exam_display):
    if not exam_display:
        return "1"
    if "第一次月考" in exam_display:
        return "1"
    if "第二次月考" in exam_display:
        return "3"
    if "月考" in exam_display:
        return "1"
    if "期中" in exam_display:
        return "2"
    if "期末" in exam_display:
        return "4"
    return "1"
