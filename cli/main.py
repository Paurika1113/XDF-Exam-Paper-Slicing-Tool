"""
XDFclier — FastAPI 后端（纯规则版）

基于 scanner.py + slice_all.py 的纯规则切片引擎。
无需 LLM，所有区域划分由 classify.py 规则完成。
"""

import io
import json
import logging
import os
import sys
import tempfile
import uuid
import zipfile
from pathlib import PurePath

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from scanner import scan_with_fallback
from slice_all import extract_type_content, generate_docx, restruct_zones

# ── 日志 ──
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    stream=sys.stdout,
)
logger = logging.getLogger("xdfclier-backend")

# ── 配置 ──
ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:3000,http://127.0.0.1:3000,http://localhost:8080,http://localhost"
)
origins = [o.strip() for o in ALLOWED_ORIGINS.split(",") if o.strip()]

TYPE_NAMES = ["论述类文本", "文学类文本", "文言文阅读", "古诗词阅读"]

# ── 响应模型 ──
class HealthResponse(BaseModel):
    status: str
    version: str = "1.0"

# ── FastAPI 应用 ──
app = FastAPI(
    title="XDFclier — 试卷切片工具",
    version="1.0",
    docs_url="/docs" if os.getenv("ENABLE_DOCS", "false").lower() == "true" else None,
    redoc_url=None,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=False,
    allow_methods=["POST", "GET"],
    allow_headers=["Content-Type", "Accept", "X-Request-ID"],
    max_age=3600,
)

@app.middleware("http")
async def add_request_id(request: Request, call_next):
    request_id = request.headers.get("X-Request-ID", str(uuid.uuid4())[:8])
    request.state.request_id = request_id
    response = await call_next(request)
    response.headers["X-Request-ID"] = request_id
    return response

# ── 接口 ──

@app.post("/api/slice")
async def slice_files(
    files: list[UploadFile] = File(...),
    grade_info: str = Form(...),
    scheme: str = Form("{}"),
    request: Request = None,
):
    """上传 DOCX → 扫描 → 规则分类 → 生成切片 DOCX → 返回 ZIP"""
    rid = getattr(request.state, "request_id", "unknown") if request else "unknown"

    try:
        scheme_dict = json.loads(scheme) if isinstance(scheme, str) else scheme
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"格式方案 JSON 无效: {exc.msg}")
    if not isinstance(scheme_dict, dict):
        raise HTTPException(status_code=400, detail="格式方案必须是 JSON 对象")

    with tempfile.TemporaryDirectory() as tmpdir:
        upload_dir = os.path.join(tmpdir, "uploads")
        output_dir = os.path.join(tmpdir, "outputs")
        os.makedirs(upload_dir)
        os.makedirs(output_dir)

        # 保存上传文件
        saved_files = []
        for f in files:
            filename = PurePath(f.filename or "").name
            if not filename.lower().endswith('.docx'):
                continue
            content = await f.read()
            path = os.path.join(upload_dir, filename)
            with open(path, "wb") as fh:
                fh.write(content)
            saved_files.append(filename)

        if not saved_files:
            raise HTTPException(status_code=400, detail="未上传有效的 .docx 文件")

        # 扫描 + 提取 + 分类（纯规则）
        from docx import Document as DocxDocument

        all_jobs = []
        for fname in saved_files:
            path = os.path.join(upload_dir, fname)
            doc = DocxDocument(path)

            ranges = scan_with_fallback(doc, fname)
            if not ranges:
                logger.warning("[%s] scan failed | file=%s", rid, fname)
                continue

            for typ, (start, end) in ranges.items():
                content = extract_type_content(doc, start, end)
                if not content:
                    continue
                from scanner import get_exam_type, get_school_name
                school = get_school_name(fname)
                exam = get_exam_type(fname)

                # 纯规则分类 zones
                zones = restruct_zones(content, typ)

                all_jobs.append({
                    "fname": fname,
                    "type": typ,
                    "school": school,
                    "exam": exam,
                    "content": content,
                    "zones": zones,
                })

        if not all_jobs:
            raise HTTPException(status_code=400, detail="未能从文件中提取任何有效内容")

        # 分组 + 生成 DOCX
        type_collections: dict[str, list] = {t: [] for t in TYPE_NAMES}
        for job in all_jobs:
            typ = job["type"]
            if typ not in type_collections:
                continue
            zones = job.get("zones", [])
            if not zones:
                zones = restruct_zones(job["content"], typ)
            type_collections[typ].append((
                job["school"], job["exam"], job["content"], zones
            ))

        type_output_names = {
            "论述类文本": f"{grade_info}论述类文本阅读切片.docx",
            "文学类文本": f"{grade_info}文学类文本阅读切片.docx",
            "文言文阅读": f"{grade_info}文言文阅读切片.docx",
            "古诗词阅读": f"{grade_info}古诗阅读切片.docx",
        }

        file_count = 0
        for typ in TYPE_NAMES:
            items = type_collections[typ]
            if not items:
                continue
            out = os.path.join(output_dir, type_output_names[typ])
            generate_docx(typ, items, out, grade_label=grade_info, format_scheme=scheme_dict)
            file_count += 1

        if file_count == 0:
            raise HTTPException(status_code=500, detail="切片生成失败")

        # 打包 ZIP
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for fn in os.listdir(output_dir):
                fp = os.path.join(output_dir, fn)
                zf.write(fp, fn)
        zip_buf.seek(0)

        logger.info("[%s] slice done | files=%d | outputs=%d", rid, len(saved_files), file_count)
        return StreamingResponse(
            zip_buf,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename=shijuan_slice_{rid}.zip"}
        )


@app.get("/api/health")
async def health():
    return {"status": "ok", "version": "1.0"}


@app.post("/api/debug-extract")
async def debug_extract(file: UploadFile = File(...)):
    if not (file.filename or '').endswith('.docx'):
        raise HTTPException(status_code=400, detail="仅支持 .docx 文件")

    try:
        from docx import Document
        content = await file.read()
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
        logger.info("debug-extract | file=%s | paragraphs=%d", file.filename, len(paragraphs))
        return {"filename": file.filename, "paragraphs": paragraphs, "count": len(paragraphs)}
    except Exception as e:
        logger.error("debug-extract error | %s", str(e))
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")


# ── debug-scan 用简化 scanner ──

TYPE_KEYWORDS_SIMPLE = {
    "论述类文本": ["阅读Ⅰ", "阅读I", "阅读1", "现代文阅读Ⅰ", "现代文阅读I",
                   "现代文阅读1", "论述类", "论述类文本", "信息类", "信息类文本",
                   "信息类文本阅读", "信息类阅读", "论述文", "实用类文本",
                   "非连续性文本", "阅读下面的文字，完成下面小题"],
    "文学类文本": ["阅读Ⅱ", "阅读II", "阅读2", "现代文阅读Ⅱ", "现代文阅读II",
                   "现代文阅读2", "文学类", "文学"],
    "文言文阅读": ["文言文阅读", "文言文", "文言知识", "课外文言文",
                   "课外文言文阅读", "阅读Ⅲ", "阅读III", "古代诗文阅读Ⅰ"],
    "古诗词阅读": ["古代诗歌阅读", "古代诗歌鉴赏", "古诗阅读",
                   "诗歌阅读", "诗歌鉴赏", "古诗词",
                   "阅读Ⅳ", "阅读IV", "古代诗文阅读Ⅱ"],
}

NON_TARGET_SIMPLE = ["名篇名句默写", "名句默写", "古诗文默写",
                     "语言文字运用", "语用", "作文", "写作",
                     "基础知识", "古代诗文阅读Ⅲ"]


def _scan_paragraphs(paras):
    l1_headers = []
    l2_headers = []
    l1_pat = __import__('re').compile(r'^[一二三四五六七八九十百]+[、.．]')
    l2_pat = __import__('re').compile(r'^[（(][一二三四五六七八九十百]+[）)]')

    for i, t in enumerate(paras):
        t = t.strip()
        if not t:
            continue
        if l1_pat.match(t):
            l1_headers.append((i, t))
        elif l2_pat.match(t):
            l2_headers.append((i, t))

    total = len(paras)

    def _next_header(idx):
        candidates = []
        for hi, ht in l1_headers:
            if hi > idx:
                candidates.append(hi)
        for hi, ht in l2_headers:
            if hi > idx:
                candidates.append(hi)
        return min(candidates) if candidates else total

    def _classify_by_keyword(text):
        for typ, keywords in TYPE_KEYWORDS_SIMPLE.items():
            for kw in keywords:
                if kw in text:
                    return typ
        return None

    result = {}
    for i, t in l2_headers:
        if any(kw in t for kw in NON_TARGET_SIMPLE):
            continue
        typ = _classify_by_keyword(t)
        if typ:
            end = _next_header(i)
            result[typ] = {"start": i, "end": end - 1, "paragraphs": paras[i:end]}

    if not result:
        for i, t in l1_headers:
            if any(kw in t for kw in NON_TARGET_SIMPLE):
                continue
            typ = _classify_by_keyword(t)
            if typ:
                end = _next_header(i)
                result[typ] = {"start": i, "end": end - 1, "paragraphs": paras[i:end]}

    return result


@app.post("/api/debug-scan")
async def debug_scan(file: UploadFile = File(...)):
    filename = PurePath(file.filename or "").name
    if not filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="仅支持 .docx 文件")

    try:
        from docx import Document
        content = await file.read()
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    ranges = scan_with_fallback(doc, filename) or {}
    result = {}
    for type_name, range_value in ranges.items():
        if isinstance(range_value, dict):
            start = int(range_value.get("start", 0))
            end = int(range_value.get("end", start))
        else:
            start, end = range_value
        start = max(0, start)
        end = min(len(paragraphs) - 1, end)
        if start <= end:
            result[type_name] = {
                "start": start,
                "end": end,
                "paragraphs": paragraphs[start:end + 1],
            }

    if not result:
        result = _scan_paragraphs(paragraphs)

    return {"filename": filename, "types": result, "total_paragraphs": len(paragraphs)}


# ── 静态文件服务（前端）- 本地开发用，Vercel 部署由 CDN 处理
# FRONTEND_DIR = os.path.join(os.path.dirname(__file__), "..", "frontend")
# if os.path.isdir(FRONTEND_DIR):
#     app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")
